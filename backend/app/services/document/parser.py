# backend/app/services/document/parser.py
"""
Document Parser
Supports: PDF, DOCX, TXT, XLSX
Returns: ParsedDocument with raw text, metadata, and page-level info
"""
import hashlib
import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import pypdf
from docx import Document as DocxDocument
import openpyxl


@dataclass
class PageContent:
    page_number: int
    text: str
    char_count: int = 0

    def __post_init__(self):
        self.char_count = len(self.text)


@dataclass
class ParsedDocument:
    title: Optional[str]
    author: Optional[str]
    pages: list[PageContent]
    full_text: str = ""
    page_count: int = 0
    word_count: int = 0
    file_hash: str = ""

    def __post_init__(self):
        self.full_text = "\n\n".join(p.text for p in self.pages)
        self.page_count = len(self.pages)
        self.word_count = len(self.full_text.split())


def compute_file_hash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def parse_pdf(content: bytes) -> ParsedDocument:
    reader = pypdf.PdfReader(io.BytesIO(content))
    meta = reader.metadata or {}
    pages = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(PageContent(page_number=i, text=text.strip()))

    return ParsedDocument(
        title=meta.get("/Title") or None,
        author=meta.get("/Author") or None,
        pages=pages,
        file_hash=compute_file_hash(content),
    )


def parse_docx(content: bytes) -> ParsedDocument:
    doc = DocxDocument(io.BytesIO(content))
    # Treat each 40 paragraphs as a "page" approximation
    full_paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    chunk_size = 40
    pages = []
    for i in range(0, max(len(full_paras), 1), chunk_size):
        chunk_text = "\n".join(full_paras[i : i + chunk_size])
        pages.append(PageContent(page_number=len(pages) + 1, text=chunk_text))

    core_props = doc.core_properties
    return ParsedDocument(
        title=core_props.title or None,
        author=core_props.author or None,
        pages=pages,
        file_hash=compute_file_hash(content),
    )


def parse_txt(content: bytes) -> ParsedDocument:
    text = content.decode("utf-8", errors="replace")
    lines = text.splitlines()
    # 100 lines per page
    chunk_size = 100
    pages = []
    for i in range(0, max(len(lines), 1), chunk_size):
        chunk_text = "\n".join(lines[i : i + chunk_size])
        pages.append(PageContent(page_number=len(pages) + 1, text=chunk_text))

    return ParsedDocument(
        title=None,
        author=None,
        pages=pages,
        file_hash=compute_file_hash(content),
    )


def parse_xlsx(content: bytes) -> ParsedDocument:
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    pages = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            rows_text.append(" | ".join(cells))
        page_text = f"Sheet: {sheet_name}\n" + "\n".join(rows_text)
        pages.append(PageContent(page_number=len(pages) + 1, text=page_text))
    wb.close()

    return ParsedDocument(
        title=None,
        author=None,
        pages=pages,
        file_hash=compute_file_hash(content),
    )


PARSERS = {
    "pdf":  parse_pdf,
    "docx": parse_docx,
    "txt":  parse_txt,
    "xlsx": parse_xlsx,
}


def parse_document(content: bytes, file_type: str) -> ParsedDocument:
    parser = PARSERS.get(file_type)
    if not parser:
        raise ValueError(f"Unsupported file type: {file_type}")
    return parser(content)
